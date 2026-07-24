"""Document generation — turn text into a downloadable file.

E.V. can produce a .txt, .md, .pdf or .docx (Word) file from any content and
send it back over Telegram. Everything is built in memory (BytesIO), so nothing
touches the disk. PDF uses reportlab and Word uses python-docx (both pure-Python,
no system libraries — safe on the tiny VM).
"""

from __future__ import annotations

import io
import re
from xml.sax.saxutils import escape

# Format aliases -> canonical extension.
_ALIASES = {
    "txt": "txt", "texto": "txt", "text": "txt",
    "md": "md", "markdown": "md",
    "pdf": "pdf",
    "docx": "docx", "doc": "docx", "word": "docx",
}

SUPPORTED = ("txt", "md", "pdf", "docx")


def normalize_format(fmt: str | None) -> str | None:
    """Map a user-supplied format word to a canonical extension (or None)."""
    if not fmt:
        return None
    return _ALIASES.get(fmt.strip().lower())


def slugify(title: str) -> str:
    slug = re.sub(r"[^\w\-]+", "_", (title or "").strip().lower()).strip("_")
    return slug[:40] or "documento"


def build(fmt: str, title: str, content: str) -> tuple[bytes, str]:
    """Build a document. Returns (file_bytes, filename).

    Raises ValueError if the format is not supported.
    """
    ext = normalize_format(fmt)
    if ext not in SUPPORTED:
        raise ValueError(
            f"Formato '{fmt}' não suportado. Use: {', '.join(SUPPORTED)} (ou 'word')."
        )
    title = (title or "Documento").strip()
    content = content or ""
    filename = f"{slugify(title)}.{ext}"
    if ext == "txt":
        return _plain(title, content), filename
    if ext == "md":
        return _markdown(title, content), filename
    if ext == "pdf":
        return _pdf(title, content), filename
    return _docx(title, content), filename  # docx


def _plain(title: str, content: str) -> bytes:
    body = f"{title}\n{'=' * len(title)}\n\n{content}\n"
    return body.encode("utf-8")


def _markdown(title: str, content: str) -> bytes:
    return f"# {title}\n\n{content}\n".encode("utf-8")


def _pdf(title: str, content: str) -> bytes:
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        topMargin=2 * cm, bottomMargin=2 * cm,
        leftMargin=2 * cm, rightMargin=2 * cm,
        title=title,
    )
    styles = getSampleStyleSheet()
    body = styles["BodyText"]
    body.alignment = TA_LEFT
    body.leading = 15
    story = [Paragraph(escape(title), styles["Title"]), Spacer(1, 14)]
    for line in content.split("\n"):
        if line.strip():
            story.append(Paragraph(escape(line), body))
        else:
            story.append(Spacer(1, 8))
    if len(story) <= 2:  # title only, no body
        story.append(Paragraph("(sem conteúdo)", body))
    doc.build(story)
    return buf.getvalue()


def _docx(title: str, content: str) -> bytes:
    from docx import Document

    d = Document()
    d.add_heading(title, level=0)
    for line in content.split("\n"):
        d.add_paragraph(line)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()
